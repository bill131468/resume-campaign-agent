"""简历导出服务 - 生成 Word 和 PDF 格式的简历文件"""

from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class ResumeExporter:
    """将简历数据导出为格式化的 Word/PDF 文档"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """设置默认样式"""
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)

    def _add_section_heading(self, text):
        """添加章节标题"""
        heading = self.doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.size = Pt(14)
        return heading

    def _add_basic_info(self, profile):
        """基本信息"""
        self._add_section_heading("基本信息")

        info_pairs = []
        if profile.get("full_name"):
            info_pairs.append(("姓名", profile["full_name"]))
        if profile.get("preferred_name"):
            info_pairs.append(("常用名", profile["preferred_name"]))
        if profile.get("phone"):
            info_pairs.append(("手机", profile["phone"]))
        if profile.get("email"):
            info_pairs.append(("邮箱", profile["email"]))
        if profile.get("city"):
            info_pairs.append(("所在城市", profile["city"]))
        if profile.get("years_experience"):
            info_pairs.append(("工作年限", f"{profile['years_experience']}年"))
        if profile.get("job_seeking_status"):
            info_pairs.append(("求职状态", profile["job_seeking_status"]))
        if profile.get("available_date"):
            info_pairs.append(("可入职日期", str(profile["available_date"])))
        if profile.get("base_locations"):
            info_pairs.append(("期望城市", "、".join(profile["base_locations"])))
        if profile.get("target_roles"):
            info_pairs.append(("期望职位", "、".join(profile["target_roles"])))

        if not info_pairs:
            return

        rows = (len(info_pairs) + 1) // 2
        table = self.doc.add_table(rows=rows, cols=4, style='Table Grid')
        table.autofit = True

        for idx, (label, value) in enumerate(info_pairs):
            row = idx // 2
            col = (idx % 2) * 2
            label_cell = table.cell(row, col)
            value_cell = table.cell(row, col + 1)
            label_cell.text = label
            value_cell.text = str(value) if value else ""
            label_cell.width = Inches(1.2)
            value_cell.width = Inches(2.2)

        self.doc.add_paragraph()

    def _add_professional_summary(self, profile):
        """职业摘要"""
        if profile.get("professional_headline"):
            self._add_section_heading("职业标题")
            self.doc.add_paragraph(profile["professional_headline"])

        if profile.get("summary"):
            self._add_section_heading("个人总结")
            self.doc.add_paragraph(profile["summary"])

        self.doc.add_paragraph()

    def _add_education(self, education_list):
        """教育经历"""
        if not education_list:
            return

        self._add_section_heading("教育经历")
        for edu in education_list:
            p = self.doc.add_paragraph()
            run_school = p.add_run(edu.get("school", ""))
            run_school.bold = True
            run_school.font.size = Pt(11)

            degree_major = f"  |  {edu.get('degree', '')} · {edu.get('major', '')}"
            if edu.get("minor"):
                degree_major += f"（辅修：{edu['minor']}）"
            p.add_run(degree_major)

            date_str = ""
            if edu.get("start_date") and edu.get("end_date"):
                date_str = f"{edu['start_date']} - {edu['end_date']}"
            elif edu.get("graduation_year"):
                date_str = f"毕业年份：{edu['graduation_year']}"
            if date_str:
                p.add_run(f"\n{date_str}").italic = True

            if edu.get("gpa") is not None:
                scale = f"/{edu['gpa_scale']}" if edu.get("gpa_scale") else ""
                p.add_run(f"  |  GPA：{edu['gpa']}{scale}")

            if edu.get("rank"):
                p.add_run(f"  |  排名：{edu['rank']}")

            if edu.get("core_courses"):
                p.add_run(f"\n核心课程：{'、'.join(edu['core_courses'])}")

        self.doc.add_paragraph()

    def _add_work_experience(self, work_list):
        """工作经历"""
        if not work_list:
            return

        self._add_section_heading("工作经历")
        for exp in work_list:
            p = self.doc.add_paragraph()
            run_company = p.add_run(exp.get("company", ""))
            run_company.bold = True
            run_company.font.size = Pt(11)
            p.add_run(f"  |  {exp.get('title', '')}")

            meta = []
            if exp.get("start_date"):
                end = str(exp["end_date"]) if exp.get("end_date") else "至今"
                meta.append(f"{exp['start_date']} - {end}")
            if exp.get("location"):
                meta.append(exp["location"])
            if exp.get("department"):
                meta.append(exp["department"])
            if meta:
                p.add_run(f"\n{'  |  '.join(meta)}").italic = True

            if exp.get("experience_type"):
                type_map = {
                    "full_time": "全职", "part_time": "兼职", "internship": "实习",
                    "contract": "合同制", "freelance": "自由职业"
                }
                p.add_run(f"  |  {type_map.get(exp['experience_type'], exp['experience_type'])}")

            if exp.get("responsibilities"):
                p_resp = self.doc.add_paragraph()
                p_resp.add_run("主要职责：").bold = True
                p_resp.add_run(exp["responsibilities"])

            if exp.get("highlights"):
                for highlight in exp["highlights"]:
                    self.doc.add_paragraph(highlight, style='List Bullet')

            self.doc.add_paragraph()

        self.doc.add_paragraph()

    def _add_projects(self, projects_list):
        """项目经历"""
        if not projects_list:
            return

        self._add_section_heading("项目经历")
        for proj in projects_list:
            p = self.doc.add_paragraph()
            run_name = p.add_run(proj.get("name", ""))
            run_name.bold = True
            run_name.font.size = Pt(11)

            if proj.get("role"):
                p.add_run(f"  |  {proj['role']}")

            if proj.get("start_date"):
                end = str(proj["end_date"]) if proj.get("end_date") else "至今"
                p.add_run(f"\n{proj['start_date']} - {end}").italic = True

            self.doc.add_paragraph(proj.get("description", ""))

            if proj.get("highlights"):
                for highlight in proj["highlights"]:
                    self.doc.add_paragraph(highlight, style='List Bullet')

            if proj.get("skills"):
                p_skills = self.doc.add_paragraph()
                p_skills.add_run("技术栈：").bold = True
                p_skills.add_run("、".join(proj["skills"]))

            self.doc.add_paragraph()

    def _add_skills(self, profile):
        """技能"""
        if profile.get("skills"):
            self._add_section_heading("专业技能")
            self.doc.add_paragraph("、".join(profile["skills"]))
            self.doc.add_paragraph()

    def _add_certificates(self, certs_list):
        """证书"""
        if not certs_list:
            return

        self._add_section_heading("证书与资质")
        for cert in certs_list:
            text = cert.get("name", "")
            if cert.get("issuer"):
                text += f"  -  {cert['issuer']}"
            if cert.get("obtained_at"):
                text += f"  ({cert['obtained_at']})"
            if cert.get("score"):
                text += f"  成绩：{cert['score']}"
            self.doc.add_paragraph(text, style='List Bullet')
        self.doc.add_paragraph()

    def _add_awards(self, awards_list):
        """获奖"""
        if not awards_list:
            return

        self._add_section_heading("获奖经历")
        for award in awards_list:
            text = award.get("name", "")
            if award.get("issuer"):
                text += f"  -  {award['issuer']}"
            if award.get("awarded_at"):
                text += f"  ({award['awarded_at']})"
            self.doc.add_paragraph(text, style='List Bullet')
            if award.get("description"):
                self.doc.add_paragraph(award["description"])
        self.doc.add_paragraph()

    def _add_languages(self, lang_details):
        """语言能力"""
        if not lang_details:
            return

        self._add_section_heading("语言能力")
        for lang in lang_details:
            text = f"{lang.get('language', '')}：{lang.get('proficiency', '')}"
            if lang.get("test_name"):
                text += f"  |  {lang['test_name']}"
            if lang.get("score"):
                text += f"：{lang['score']}"
            self.doc.add_paragraph(text, style='List Bullet')
        self.doc.add_paragraph()

    def _add_campus(self, campus_list):
        """校园经历"""
        if not campus_list:
            return

        self._add_section_heading("校园经历")
        for exp in campus_list:
            p = self.doc.add_paragraph()
            run_org = p.add_run(exp.get("organization", ""))
            run_org.bold = True
            p.add_run(f"  |  {exp.get('role', '')}")
            if exp.get("start_date"):
                end = str(exp["end_date"]) if exp.get("end_date") else "至今"
                p.add_run(f"\n{exp['start_date']} - {end}").italic = True
            self.doc.add_paragraph(exp.get("description", ""))
        self.doc.add_paragraph()

    def _add_self_evaluation(self, profile):
        """自我评价"""
        if profile.get("self_evaluation"):
            self._add_section_heading("自我评价")
            self.doc.add_paragraph(profile["self_evaluation"])
            self.doc.add_paragraph()

    def export(self, profile: dict, company: str = "", position: str = "") -> BytesIO:
        """导出简历为 Word 文档"""
        self.doc = Document()
        self._setup_styles()

        full_name = profile.get("full_name") or profile.get("preferred_name") or "个人简历"
        title_text = full_name
        if company or position:
            parts = [full_name]
            if company:
                parts.append(company)
            if position:
                parts.append(position)
            title_text = " - ".join(parts)
        title_text += " · 个人简历"

        title_heading = self.doc.add_heading(title_text, level=0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

        self._add_basic_info(profile)
        self._add_professional_summary(profile)
        self._add_education(profile.get("education", []))
        self._add_work_experience(profile.get("work_experience", []))
        self._add_projects(profile.get("projects", []))
        self._add_skills(profile)
        self._add_certificates(profile.get("certificates", []))
        self._add_awards(profile.get("awards", []))
        self._add_languages(profile.get("language_details", []))
        self._add_campus(profile.get("campus_experience", []))
        self._add_self_evaluation(profile)

        self.doc.core_properties.title = f"{full_name} 个人简历"
        self.doc.core_properties.author = full_name
        self.doc.core_properties.subject = "个人简历"
        self.doc.core_properties.keywords = "简历,求职,个人简历"
        self.doc.core_properties.comments = "由 Resume Campaign Agent 自动生成"
        self.doc.core_properties.category = "简历"
        self.doc.core_properties.language = "zh-CN"
        self.doc.core_properties.identifier = f"resume-{full_name}"
        self.doc.core_properties.created = datetime.now()
        self.doc.core_properties.modified = datetime.now()

        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _build_resume_html(self, profile: dict, full_name: str) -> str:
        """构建简历 HTML"""
        html = f"""<html><head><meta charset="utf-8"><style>
body {{ font-family: 'Microsoft YaHei', sans-serif; font-size: 14px; margin: 40px; }}
h1 {{ text-align: center; color: #003366; font-size: 24px; }}
h2 {{ color: #003366; border-bottom: 1px solid #ccc; padding-bottom: 4px; font-size: 18px; }}
table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
td {{ border: 1px solid #ccc; padding: 8px; }}
p {{ line-height: 1.6; }}
</style></head><body>
<h1>{full_name} - 个人简历</h1>
<h2>基本信息</h2>
<table>
<tr><td><b>姓名</b></td><td>{profile.get("full_name", "")}</td><td><b>手机</b></td><td>{profile.get("phone", "")}</td></tr>
<tr><td><b>邮箱</b></td><td>{profile.get("email", "")}</td><td><b>城市</b></td><td>{profile.get("city", "")}</td></tr>
</table>
"""
        if profile.get("summary"):
            html += f"<h2>个人总结</h2><p>{profile['summary']}</p>"

        if profile.get("education"):
            html += "<h2>教育经历</h2>"
            for edu in profile["education"]:
                html += f"<p><b>{edu.get('school', '')}</b> | {edu.get('degree', '')} · {edu.get('major', '')} | {edu.get('graduation_year', '')}</p>"

        if profile.get("work_experience"):
            html += "<h2>工作经历</h2>"
            for exp in profile["work_experience"]:
                html += f"<p><b>{exp.get('company', '')}</b> | {exp.get('title', '')}</p>"
                if exp.get("responsibilities"):
                    html += f"<p>{exp['responsibilities']}</p>"
                if exp.get("highlights"):
                    for h in exp["highlights"]:
                        html += f"<p>• {h}</p>"

        if profile.get("projects"):
            html += "<h2>项目经历</h2>"
            for proj in profile["projects"]:
                html += f"<p><b>{proj.get('name', '')}</b> | {proj.get('role', '')}</p>"
                html += f"<p>{proj.get('description', '')}</p>"

        if profile.get("skills"):
            html += "<h2>专业技能</h2>"
            html += f"<p>{'、'.join(profile['skills'])}</p>"

        if profile.get("certificates"):
            html += "<h2>证书与资质</h2>"
            for cert in profile["certificates"]:
                html += f"<p>{cert.get('name', '')}</p>"

        if profile.get("self_evaluation"):
            html += f"<h2>自我评价</h2><p>{profile['self_evaluation']}</p>"

        html += "</body></html>"
        return html

    async def export_html_to_pdf(self, profile: dict, company: str = "", position: str = "") -> BytesIO:
        """生成 PDF 简历（腾讯等平台兼容）"""
        full_name = profile.get("full_name") or profile.get("preferred_name") or "个人简历"
        html = self._build_resume_html(profile, full_name)

        from playwright.async_api import async_playwright
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            await page.set_content(html, wait_until="networkidle")
            pdf_bytes = await page.pdf(format="A4")
            await browser.close()

        buffer = BytesIO(pdf_bytes)
        buffer.seek(0)
        return buffer